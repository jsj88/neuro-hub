"""Generate detailed DOCX report for LME T-maze pipeline results."""

import os
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

RESULTS_DIR = os.path.dirname(os.path.abspath(__file__))
FIGURE_PATH = os.path.join(RESULTS_DIR, "12_lme_tmaze_pipeline.png")
FIGURE2_PATH = os.path.join(RESULTS_DIR, "12_lme_cdi_rtdiff.png")

# Load data
lme = pd.read_csv(os.path.join(RESULTS_DIR, "lme_all_results.csv"))
info = pd.read_csv(os.path.join(RESULTS_DIR, "subject_info.csv"))
inter = pd.read_csv(os.path.join(RESULTS_DIR, "interaction_models.csv"))
rtdiff = pd.read_csv(os.path.join(RESULTS_DIR, "rtdiff_lme_results.csv"))
trial_df = pd.read_csv(os.path.join(RESULTS_DIR, "trial_data_with_rpes.csv"))


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Set font size for entire table
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Calibri"

    return table


# =====================================================================
# BUILD DOCUMENT
# =====================================================================
doc = Document()

# -- Styles --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

# =====================================================================
# TITLE PAGE
# =====================================================================
title = doc.add_heading("Linear Mixed-Effects Analysis of Neural-Behavioral "
                        "Relationships in the T-Maze AR Task", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Analysis Report")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Dataset: merged_clean_V3_scaled.csv").font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("N = 43 participants | 7,963 single trials | 13 EEG channels").font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Model: ERP ~ Behavioral Predictor + (1 | Subject) + (1 | Trial)")
run.bold = True
run.font.size = Pt(11)

doc.add_page_break()

# =====================================================================
# 1. OVERVIEW
# =====================================================================
doc.add_heading("1. Overview", level=1)

doc.add_paragraph(
    "This report describes a comprehensive linear mixed-effects (LME) analysis "
    "examining the relationship between behavioral model-derived variables and "
    "single-trial ERP amplitudes in a T-maze augmented reality (AR) task. "
    "Participants navigated a virtual T-maze using HoloLens 2, making left/right "
    "choices with probabilistic reward feedback. ERP amplitudes were extracted "
    "at 13 scalp sites. Behavioral predictors were derived from computational "
    "models of reward learning fit to each participant's choice data."
)

doc.add_heading("1.1 Motivation", level=2)
doc.add_paragraph(
    "Previous analyses of this dataset used within-subject Pearson correlations "
    "between model-derived reward prediction errors (RPEs) and ERP amplitudes, "
    "followed by Fisher-z transformation and group-level one-sample t-tests. "
    "This approach has two key limitations: (1) it discards trial-level information "
    "by summarizing each subject as a single correlation coefficient, and (2) it "
    "does not account for the hierarchical structure of the data (trials nested "
    "within subjects). Linear mixed-effects models address both limitations by "
    "modeling all single trials simultaneously while accounting for random "
    "variability across subjects and trial positions."
)

doc.add_heading("1.2 Statistical Model", level=2)
doc.add_paragraph(
    "All analyses use the following LME specification:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ERP_amplitude ~ Behavioral_predictor + (1 | Subject) + (1 | Trial)")
run.bold = True
run.italic = True

doc.add_paragraph("Where:")
doc.add_paragraph(
    "The fixed effect (Behavioral_predictor) estimates the population-level "
    "relationship between the predictor and ERP amplitude.",
    style="List Bullet"
)
doc.add_paragraph(
    "(1 | Subject) is a random intercept for each participant, accounting for "
    "between-subject differences in overall ERP baseline amplitude.",
    style="List Bullet"
)
doc.add_paragraph(
    "(1 | Trial) is a variance component for trial position, accounting for "
    "shared trial-order effects (e.g., learning, fatigue) across participants. "
    "Implemented via statsmodels vc_formula with C(Trial) as a variance component.",
    style="List Bullet"
)
doc.add_paragraph(
    "Models were estimated using restricted maximum likelihood (REML) with the "
    "Powell optimizer in Python's statsmodels (v0.14.6). For models where the "
    "trial variance component caused convergence issues, a fallback to "
    "(1 | Subject) only was used."
)

doc.add_page_break()

# =====================================================================
# 2. DATA AND VARIABLES
# =====================================================================
doc.add_heading("2. Data and Variables", level=1)

doc.add_heading("2.1 Participants and Task", level=2)
doc.add_paragraph(
    f"The dataset contains {len(info)} participants who completed the T-maze AR "
    f"task. Participants performed an average of {info['n_trials'].mean():.1f} trials "
    f"(SD = {info['n_trials'].std():.1f}, range = {info['n_trials'].min()}-"
    f"{info['n_trials'].max()}), yielding 7,963 total trials. On each trial, "
    "participants walked to one of two choice locations in augmented reality and "
    "received probabilistic reward feedback."
)

doc.add_heading("2.2 Reward Conditions", level=2)
doc.add_paragraph(
    "Reward conditions were defined by walking speed (Vel_Cond):"
)

add_table(doc,
    ["Vel_Cond", "Label", "Earned", "N Trials", "Description"],
    [
        ["1", "High Reward", "$0.25", "1,580", "Slow walking RT -> high reward"],
        ["2", "No Reward", "$0.00", "4,011", "No reward regardless of RT"],
        ["3", "Low Reward", "$0.05", "2,372", "Fast/average walking RT -> low reward"],
    ]
)

doc.add_paragraph("")

doc.add_heading("2.3 Behavioral Model Fitting", level=2)
doc.add_paragraph(
    "Ten computational models of reward learning were fit to each participant's "
    "binary choice and outcome data using maximum likelihood estimation with "
    "differential evolution optimization: Random, Win-Stay-Lose-Shift (WSLS), "
    "Rescorla-Wagner (RW), Q-learning, Q-dual, Actor-Critic, Q-decay, "
    "RW+Bias, Choice Kernel (CK), and RW+CK."
)

doc.add_heading("Best-Fitting Model Distribution", level=3)
model_counts = info["best_model"].value_counts()
rows = []
for model_name, count in model_counts.items():
    rows.append([model_name, str(count), f"{count/len(info)*100:.1f}%"])
add_table(doc, ["Model", "N Subjects", "Percentage"], rows)

doc.add_paragraph("")
doc.add_paragraph(
    f"The Rescorla-Wagner model was fit to all participants to extract trial-level "
    f"latent variables. Two versions were fit: (1) with binary outcomes (reward=1, "
    f"no-reward=0), yielding binary RPEs, and (2) with magnitude outcomes "
    f"(Earned: $0.00, $0.05, $0.25), yielding magnitude RPEs."
)

doc.add_paragraph(
    f"Binary RW learning rate: M = {info['rw_alpha_binary'].mean():.3f}, "
    f"SD = {info['rw_alpha_binary'].std():.3f}. "
    f"Magnitude RW learning rate: M = {info['rw_alpha_mag'].mean():.3f}, "
    f"SD = {info['rw_alpha_mag'].std():.3f}."
)

doc.add_heading("2.4 Behavioral Predictors", level=2)
doc.add_paragraph(
    "Eleven behavioral predictors were tested in the LME models:"
)
add_table(doc,
    ["Predictor", "Description", "Source"],
    [
        ["rpe_binary", "Reward prediction error (binary outcome)", "RW model fit to 0/1 outcomes"],
        ["rpe_mag", "Reward prediction error (magnitude outcome)", "RW model fit to $0.00-$0.25 outcomes"],
        ["abs_rpe_binary", "Unsigned (absolute) RPE, binary", "|rpe_binary|"],
        ["abs_rpe_mag", "Unsigned (absolute) RPE, magnitude", "|rpe_mag|"],
        ["chosen_q_binary", "Q-value of chosen option (binary)", "Q(chosen) from binary RW model"],
        ["chosen_q_mag", "Q-value of chosen option (magnitude)", "Q(chosen) from magnitude RW model"],
        ["is_reward", "Binary reward indicator", "1 = reward, 0 = no reward"],
        ["reward_level_num", "Ordinal reward level", "0 = none, 1 = low ($0.05), 2 = high ($0.25)"],
        ["Earned", "Dollar amount earned on trial", "$0.00, $0.05, or $0.25"],
        ["Start_RT", "Walking reaction time (seconds)", "Time from countdown to feedback"],
        ["RT_DIFF", "RT change (prev - current Start_RT)", "Positive = sped up, Negative = slowed down"],
    ]
)

doc.add_paragraph("")

doc.add_heading("2.5 ERP Channels", level=2)
doc.add_paragraph(
    "ERP amplitudes were available for 13 scalp sites: FCz, Cz, FC1, FC2, "
    "F4, F5, C4, C3, P4, P3, P8, P7, and Pz. These represent the scaled "
    "mean amplitude in the feedback-locked ERP waveform. FCz and Cz are the "
    "primary channels of interest for the reward positivity (RewP) / "
    "feedback-related negativity (FRN) component."
)

doc.add_heading("2.6 Conditions", level=2)
doc.add_paragraph(
    "Each predictor-channel combination was tested under eleven data subsets:"
)
add_table(doc,
    ["Condition", "Definition", "N Trials"],
    [
        ["all", "All trials", "7,963"],
        ["high_rew", "Vel_Cond = 1 (High Reward, $0.25)", "1,580"],
        ["low_rew", "Vel_Cond = 3 (Low Reward, $0.05)", "2,372"],
        ["no_rew", "Vel_Cond = 2 (No Reward, $0.00)", "4,011"],
        ["any_rew", "Reward = 1 (any reward trial)", "~3,952"],
        ["stay", "PES_PRS_index = 1 (stayed at same location)", "~5,000"],
        ["shift", "PES_PRS_index = 2 (shifted to other location)", "~2,700"],
        ["rew_stay", "CDI = 1: Reward + Stay", "~2,035"],
        ["rew_shift", "CDI = 2: Reward + Shift", "~1,900"],
        ["norew_stay", "CDI = 3: No-Reward + Stay", "~1,473"],
        ["norew_shift", "CDI = 4: No-Reward + Shift", "~2,512"],
    ]
)
doc.add_paragraph("")
doc.add_paragraph(
    "The four CDI (Condition_Direction_index) conditions cross reward outcome "
    "with stay/shift behavior, enabling analysis of how RPE-ERP coupling varies "
    "across combinations of reward outcome and behavioral strategy."
)

doc.add_page_break()

# =====================================================================
# 3. ANALYSIS STEPS
# =====================================================================
doc.add_heading("3. Analysis Pipeline", level=1)

doc.add_heading("Step 1: Data Loading and Preparation", level=2)
doc.add_paragraph(
    "The behavioral CSV (merged_clean_V3_scaled.csv) was loaded. Variables were "
    "recoded: Direction (1=Left, 2=Right) was converted to choice (0=Left, 1=Right). "
    "Reward (1=reward, 2=no-reward) was converted to binary outcome (1/0). "
    "Vel_Cond was mapped to reward level labels. PES_PRS_index was mapped to "
    "stay/shift labels. Condition_Direction_index was mapped to combined "
    "outcome-by-direction labels."
)

doc.add_heading("Step 2: Computational Model Fitting", level=2)
doc.add_paragraph(
    "For each participant, 10 reward learning models were fit to binary choice "
    "and outcome data using maximum likelihood with differential evolution. "
    "The best-fitting model was identified by BIC. Separately, the Rescorla-Wagner "
    "(RW) model was fit to both binary and magnitude outcomes. Trial-level latent "
    "variables were extracted: RPEs (reward prediction errors), |RPE| (unsigned RPEs), "
    "and Q-values of the chosen option."
)

doc.add_heading("Step 3: Mass Univariate LME Analysis", level=2)
doc.add_paragraph(
    f"A total of {len(lme)} LME models were fit, one for each combination of: "
    "11 behavioral predictors x 13 ERP channels x 11 conditions "
    "(not all combinations are possible for predictors like is_reward that are "
    "constant within some condition subsets). Each model had the form:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Channel_amplitude ~ Predictor + (1 | Subject) + (1 | Trial)")
run.italic = True

doc.add_paragraph(
    "The (1 | Trial) variance component was implemented using statsmodels' "
    "vc_formula parameter with C(Trial) as a categorical variance component. "
    "If convergence failed with both random effects, the model fell back to "
    "(1 | Subject) only."
)

doc.add_heading("Step 4: Multiple Comparison Correction", level=2)
doc.add_paragraph(
    f"Benjamini-Hochberg false discovery rate (FDR) correction was applied "
    f"across all {len(lme)} ERP tests at alpha = 0.05, and separately across "
    f"all {len(rtdiff)} RT_DIFF tests. This controls the expected proportion "
    f"of false discoveries among rejected hypotheses."
)

doc.add_heading("Step 5: Interaction Models", level=2)
doc.add_paragraph(
    "Separate interaction LME models were fit to test whether the RPE-ERP "
    "relationship was moderated by experimental factors:"
)
doc.add_paragraph(
    "RPE x Reward Level: ERP ~ RPE * C(reward_level) + (1 | Subject), "
    "with no-reward as the reference category.",
    style="List Bullet"
)
doc.add_paragraph(
    "RPE x Stay/Shift: ERP ~ RPE * C(PES_PRS_index) + (1 | Subject), "
    "testing whether the RPE-ERP coupling differs for stay vs. shift trials.",
    style="List Bullet"
)
doc.add_paragraph(
    "RPE x Velocity Condition: ERP ~ RPE * C(Vel_Cond) + (1 | Subject), "
    "testing moderation by the three velocity/reward conditions.",
    style="List Bullet"
)
doc.add_paragraph(
    "RPE x Condition_Direction_index (CDI): ERP ~ RPE * C(cond_dir) + (1 | Subject), "
    "testing whether RPE-ERP coupling differs across reward-stay, reward-shift, "
    "no-reward-stay, and no-reward-shift trials. Reference category: norew_stay.",
    style="List Bullet"
)
doc.add_paragraph(
    "RPE x CDI -> RT_DIFF: RT_DIFF ~ RPE * C(cond_dir) + (1 | Subject), "
    "testing whether RPE predicts RT adjustment differently across CDI conditions.",
    style="List Bullet"
)

doc.add_heading("Step 6: Post-Outcome Behavioral Adjustment", level=2)
doc.add_paragraph(
    "To examine whether RPEs influence subsequent behavior, LME models tested "
    "whether the RPE on trial t-1 predicted the change in walking reaction time "
    "on trial t: RT_change ~ prev_RPE + (1 | Subject)."
)

doc.add_heading("Step 6b: RT_DIFF as Outcome Variable", level=2)
doc.add_paragraph(
    "RT_DIFF was defined as the previous trial's walking RT (Start_RT) minus "
    "the current trial's walking RT. Positive values indicate speeding up "
    "(shorter current RT), negative values indicate slowing down. "
    "For each behavioral predictor and condition, we fit:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RT_DIFF ~ Predictor + (1 | Subject) + (1 | Trial)")
run.italic = True
doc.add_paragraph(
    f"This yielded {len(rtdiff)} RT_DIFF LME models, testing which behavioral "
    "variables predict trial-to-trial walking speed adjustment across all "
    "11 condition subsets."
)

doc.add_page_break()

# =====================================================================
# 4. RESULTS
# =====================================================================
doc.add_heading("4. Results", level=1)

# 4.1 Overall summary
doc.add_heading("4.1 Overall Model Summary", level=2)
n_sig_unc = (lme["p"] < 0.05).sum()
n_sig_fdr = lme["sig_fdr"].sum()
doc.add_paragraph(
    f"Of {len(lme)} ERP LME models fit, {n_sig_unc} ({n_sig_unc/len(lme)*100:.1f}%) "
    f"were significant at the uncorrected p < .05 level, and {n_sig_fdr} "
    f"({n_sig_fdr/len(lme)*100:.1f}%) survived FDR correction at q < .05."
)
n_sig_unc_rt = (rtdiff["p"] < 0.05).sum()
n_sig_fdr_rt = rtdiff["sig_fdr"].sum()
doc.add_paragraph(
    f"Of {len(rtdiff)} RT_DIFF LME models fit, {n_sig_unc_rt} "
    f"({n_sig_unc_rt/len(rtdiff)*100:.1f}%) were significant at p < .05, and "
    f"{n_sig_fdr_rt} ({n_sig_fdr_rt/len(rtdiff)*100:.1f}%) survived FDR correction."
)

# 4.2 Per-predictor summary
doc.add_heading("4.2 Results by Predictor", level=2)
doc.add_paragraph(
    "The following table summarizes the number of significant effects for each "
    "predictor across all channels and conditions:"
)

pred_rows = []
for pred in ["rpe_binary", "rpe_mag", "abs_rpe_binary", "abs_rpe_mag",
             "chosen_q_binary", "chosen_q_mag", "is_reward",
             "reward_level_num", "Earned", "Start_RT", "RT_DIFF"]:
    sub = lme[lme["predictor"] == pred]
    if len(sub) == 0:
        continue
    n_tot = len(sub)
    n_unc = (sub["p"] < 0.05).sum()
    n_fdr = sub["sig_fdr"].sum()
    pred_rows.append([pred, str(n_tot), str(n_unc),
                      f"{n_unc/n_tot*100:.1f}%", str(n_fdr)])

add_table(doc,
    ["Predictor", "Total Models", "p < .05", "% Sig", "FDR q < .05"],
    pred_rows
)

doc.add_paragraph("")
doc.add_paragraph(
    "Magnitude-based predictors (rpe_mag, abs_rpe_mag, chosen_q_mag) were far "
    "more predictive of ERP amplitude than their binary counterparts. The "
    "magnitude RPE yielded 49 uncorrected significant effects (41 FDR-surviving) "
    "compared to only 8 (2 FDR) for binary RPE. This suggests that the graded "
    "reward information ($0.00, $0.05, $0.25) is encoded in the ERP signal "
    "more than the simple reward/no-reward distinction."
)

# 4.3 Per-channel summary
doc.add_heading("4.3 Results by Channel", level=2)
ch_rows = []
for ch in ["FCz", "Cz", "FC1", "FC2", "F4", "F5", "C4", "C3",
           "P4", "P3", "P8", "P7", "Pz"]:
    sub = lme[lme["channel"] == ch]
    n_unc = (sub["p"] < 0.05).sum()
    n_fdr = sub["sig_fdr"].sum()
    ch_rows.append([ch, str(len(sub)), str(n_unc), str(n_fdr)])

add_table(doc,
    ["Channel", "Total", "p < .05", "FDR q < .05"],
    ch_rows
)

doc.add_paragraph("")
doc.add_paragraph(
    "Effects were broadly distributed across the scalp. FCz (the traditional "
    "RewP/FRN site) showed 21 uncorrected significant effects (19 FDR). "
    "Notably, C4 and P3 showed equal or greater numbers of significant effects, "
    "suggesting the neural correlates of reward prediction extend beyond the "
    "classic frontocentral distribution."
)

# 4.4 Per-condition summary
doc.add_heading("4.4 Results by Condition", level=2)
cond_rows = []
for cond in ["all", "high_rew", "low_rew", "no_rew", "any_rew", "stay", "shift",
             "rew_stay", "rew_shift", "norew_stay", "norew_shift"]:
    sub = lme[lme["condition"] == cond]
    n_unc = (sub["p"] < 0.05).sum()
    n_fdr = sub["sig_fdr"].sum()
    cond_rows.append([cond, str(len(sub)), str(n_unc), str(n_fdr)])

add_table(doc,
    ["Condition", "Total", "p < .05", "FDR q < .05"],
    cond_rows
)

doc.add_paragraph("")
doc.add_paragraph(
    "The 'stay' condition produced the most significant effects (49 uncorrected, "
    "35 FDR), followed by 'any_rew' (43 uncorrected, 23 FDR) and 'shift' "
    "(39 uncorrected, 23 FDR). This suggests that behavioral predictors are "
    "most strongly reflected in ERPs when participants repeat their previous "
    "choice location."
)

doc.add_page_break()

# 4.5 Key FCz results
doc.add_heading("4.5 Key Results at FCz (Primary Channel)", level=2)

doc.add_paragraph(
    "FCz is the primary site for the reward positivity (RewP), the ERP component "
    "most theoretically linked to RPE. The following table shows all significant "
    "LME effects at FCz:"
)

fcz_sig = lme[(lme["channel"] == "FCz") & (lme["p"] < 0.05)].sort_values("p")
fcz_rows = []
for _, r in fcz_sig.iterrows():
    sig_label = "***" if r["p"] < 0.001 else ("**" if r["p"] < 0.01 else "*")
    fcz_rows.append([
        r["predictor"], r["condition"],
        f"{r['estimate']:.4f}", f"{r['se']:.4f}",
        f"{r['z']:.3f}", f"{r['p']:.6f}", sig_label,
        str(int(r["n_obs"])), str(int(r["n_subj"])),
    ])

add_table(doc,
    ["Predictor", "Condition", "b", "SE", "z", "p", "Sig", "N obs", "N subj"],
    fcz_rows
)

doc.add_paragraph("")
doc.add_paragraph(
    "The strongest effects at FCz were for magnitude-based predictors. "
    "Magnitude RPE (rpe_mag) was significantly negatively associated with FCz "
    "amplitude across all conditions (all trials: b = -2.040, z = -6.11, "
    "p < .0001). The negative coefficient indicates that larger positive RPEs "
    "(better-than-expected outcomes) were associated with more negative FCz "
    "amplitudes, or equivalently, larger negative RPEs (worse-than-expected) "
    "were associated with more positive FCz amplitudes."
)

doc.add_paragraph(
    "The chosen Q-value (chosen_q_mag) showed the complementary positive "
    "relationship (all trials: b = 2.629, z = 7.17, p < .0001), indicating "
    "that higher expected value was associated with more positive FCz amplitudes."
)

doc.add_paragraph(
    "Binary RPE (rpe_binary) was not significantly associated with FCz "
    "in any condition, suggesting that the scalar reward magnitude information "
    "drives the ERP effect rather than the binary reward/no-reward distinction."
)

# 4.6 Key Cz results
doc.add_heading("4.6 Key Results at Cz", level=2)

cz_sig = lme[(lme["channel"] == "Cz") & (lme["p"] < 0.05)].sort_values("p")
cz_rows = []
for _, r in cz_sig.iterrows():
    sig_label = "***" if r["p"] < 0.001 else ("**" if r["p"] < 0.01 else "*")
    cz_rows.append([
        r["predictor"], r["condition"],
        f"{r['estimate']:.4f}", f"{r['se']:.4f}",
        f"{r['z']:.3f}", f"{r['p']:.6f}", sig_label,
    ])

add_table(doc,
    ["Predictor", "Condition", "b", "SE", "z", "p", "Sig"],
    cz_rows
)

doc.add_paragraph("")
doc.add_paragraph(
    "At Cz, significant effects included reward_level_num in the any_rew "
    "condition (b = 0.423, z = 29.84, p < .0001), Earned for any_rew trials "
    "(b = 2.113, z = 2.66, p = .008), Start_RT across all trials "
    "(b = 0.082, z = 2.64, p = .008), and magnitude RPE for any_rew "
    "(b = 1.062, z = 2.20, p = .028)."
)

doc.add_page_break()

# 4.7 Interaction models
doc.add_heading("4.7 Interaction Model Results", level=2)

doc.add_heading("4.7.1 RPE x Reward Level", level=3)
doc.add_paragraph(
    "Binary RPE x Reward Level interaction was not significant at either "
    "FCz or Cz (all interaction p > .56), indicating that the (non-significant) "
    "binary RPE-ERP relationship did not differ across reward levels."
)
doc.add_paragraph(
    "Magnitude RPE showed a significant main effect at FCz (b = -1.565, "
    "z = -2.36, p = .018) with no-reward as reference, but the interaction "
    "terms were not significant (RPE x High Reward: p = .520; RPE x Low Reward: "
    "p = .919). At Cz, there was a trending RPE x High Reward interaction "
    "(b = 1.522, z = 1.78, p = .076), suggesting the magnitude RPE-Cz "
    "relationship may be stronger for high-reward trials."
)

doc.add_heading("4.7.2 RPE x Stay/Shift", level=3)
doc.add_paragraph(
    "For binary RPE at Cz, there was a trending RPE x Stay/Shift interaction "
    "(b = 0.335, z = 1.82, p = .069), suggesting that the RPE-Cz coupling "
    "may differ between stay and shift trials. All other RPE x Stay/Shift "
    "interactions were non-significant."
)

doc.add_heading("4.7.3 RPE x Velocity Condition", level=3)
doc.add_paragraph(
    "No significant RPE x Velocity Condition interactions were observed at "
    "FCz or Cz for either binary or magnitude RPE. The magnitude RPE x "
    "Vel_Cond=2 (no reward) interaction was trending at Cz (b = -1.522, "
    "z = -1.78, p = .076)."
)

# 4.8 Post-outcome behavioral adjustment
doc.add_heading("4.8 Post-Outcome Behavioral Adjustment", level=2)

doc.add_paragraph(
    "LME models tested whether RPE on trial t-1 predicted the change in walking "
    "reaction time on trial t:"
)

add_table(doc,
    ["Predictor (t-1)", "b", "z", "p", "Significant"],
    [
        ["rpe_binary", "-0.0819", "-3.038", "0.0024", "Yes"],
        ["rpe_mag", "-0.7973", "nan", "nan", "Convergence issue"],
        ["|rpe_binary|", "-0.0019", "-0.040", "0.968", "No"],
        ["|rpe_mag|", "-0.4332", "-3.765", "0.0002", "Yes"],
    ]
)

doc.add_paragraph("")
doc.add_paragraph(
    "Binary RPE on trial t-1 significantly predicted RT change on trial t "
    "(b = -0.082, z = -3.04, p = .002): larger positive RPEs (rewarded outcomes) "
    "were followed by faster walking on the next trial, consistent with "
    "reward-driven behavioral invigoration."
)
doc.add_paragraph(
    "Unsigned magnitude RPE (|rpe_mag|) also significantly predicted RT change "
    "(b = -0.433, z = -3.77, p = .0002): larger prediction errors of either sign "
    "were followed by faster walking, suggesting that surprise (unsigned PE) "
    "triggers behavioral adjustment."
)

doc.add_heading("Walking RT Change by Previous Reward Level", level=3)
add_table(doc,
    ["Previous Trial", "Mean RT Change (s)", "Interpretation"],
    [
        ["High Reward", "-0.612", "Faster after high reward"],
        ["Low Reward", "+0.327", "Slower after low reward"],
        ["No Reward", "+0.026", "Minimal change after no reward"],
    ]
)

doc.add_paragraph("")
doc.add_paragraph(
    "Participants walked substantially faster after high-reward trials "
    "(-0.61s) and slower after low-reward trials (+0.33s), consistent with "
    "reward-dependent post-outcome behavioral invigoration/deceleration."
)

doc.add_page_break()

# =====================================================================
# 4.9 RT_DIFF as Outcome
# =====================================================================
doc.add_heading("4.9 RT_DIFF as Outcome Variable", level=2)

doc.add_paragraph(
    f"RT_DIFF (previous trial's Start_RT minus current trial's Start_RT) was "
    f"analyzed as an outcome variable to identify which behavioral predictors "
    f"drive trial-to-trial walking speed adjustments. A total of {len(rtdiff)} "
    f"LME models were fit across 10 behavioral predictors and 11 conditions."
)

doc.add_heading("4.9.1 Overall RT_DIFF Results", level=3)
n_sig_unc_rt = (rtdiff["p"] < 0.05).sum()
n_sig_fdr_rt = rtdiff["sig_fdr"].sum()
doc.add_paragraph(
    f"Of {len(rtdiff)} RT_DIFF models, {n_sig_unc_rt} ({n_sig_unc_rt/len(rtdiff)*100:.1f}%) "
    f"were significant uncorrected, and {n_sig_fdr_rt} ({n_sig_fdr_rt/len(rtdiff)*100:.1f}%) "
    f"survived FDR correction."
)

# Top RT_DIFF effects table
doc.add_heading("4.9.2 Strongest RT_DIFF Predictors", level=3)
doc.add_paragraph(
    "The following table shows the most significant RT_DIFF predictors "
    "(FDR-corrected, top 20):"
)

rt_sig = rtdiff[rtdiff["sig_fdr"]].sort_values("p").head(20)
rt_rows = []
for _, r in rt_sig.iterrows():
    sig_label = "***" if r["p"] < 0.001 else ("**" if r["p"] < 0.01 else "*")
    rt_rows.append([
        r["predictor"], r["condition"],
        f"{r['estimate']:.4f}", f"{r['se']:.4f}",
        f"{r['z']:.3f}", f"{r['p']:.6f}", sig_label,
    ])

add_table(doc,
    ["Predictor", "Condition", "b", "SE", "z", "p", "Sig"],
    rt_rows
)

doc.add_paragraph("")
doc.add_paragraph(
    "Start_RT was the strongest predictor of RT_DIFF across all conditions "
    "(b = -0.58 for all trials), reflecting regression to the mean: trials "
    "with long RTs are followed by shorter RTs and vice versa. "
    "Magnitude RPE was the second strongest behavioral predictor "
    "(b = -1.42 for all trials, b = -2.43 for any_rew trials), indicating "
    "that larger positive RPEs (better-than-expected outcomes) were followed "
    "by greater speeding up on the next trial."
)
doc.add_paragraph(
    "Earned amount and reward_level_num also significantly predicted RT_DIFF "
    "with negative coefficients, confirming that higher rewards drive "
    "behavioral invigoration (faster walking on the subsequent trial). "
    "Binary RPE predicted RT_DIFF in specific conditions (rew_stay: "
    "b = 0.12, p < .001), suggesting that within reward-stay trials, "
    "positive prediction errors led to slowing (perhaps reflecting satiation "
    "or reduced urgency)."
)

doc.add_heading("4.9.3 RT_DIFF by CDI Condition", level=3)
doc.add_paragraph(
    "Descriptive statistics for RT_DIFF across Condition_Direction_index:"
)

cdi_stats = []
for cdi in ["rew_stay", "rew_shift", "norew_stay", "norew_shift"]:
    cdi_data = trial_df[trial_df["cond_dir"] == cdi]["RT_DIFF"].dropna()
    cdi_stats.append([
        cdi, f"{cdi_data.mean():+.4f}", f"{cdi_data.std():.4f}", str(len(cdi_data))
    ])
add_table(doc,
    ["CDI Condition", "Mean RT_DIFF (s)", "SD", "N Trials"],
    cdi_stats
)

doc.add_paragraph("")
doc.add_paragraph(
    "RT_DIFF was slightly positive for reward-stay trials (+0.031s, indicating "
    "minor speeding up), near zero for reward-shift and no-reward-shift, and "
    "slightly negative for no-reward-stay (-0.008s, indicating minor slowing). "
    "The standard deviations were large (~1.8-2.0s), reflecting substantial "
    "trial-to-trial variability in walking speed."
)

doc.add_page_break()

# =====================================================================
# 4.10 CDI-Specific Results
# =====================================================================
doc.add_heading("4.10 CDI-Specific ERP Results", level=2)

doc.add_paragraph(
    "The four Condition_Direction_index conditions (rew_stay, rew_shift, "
    "norew_stay, norew_shift) allow examination of how reward outcome and "
    "behavioral strategy jointly modulate the RPE-ERP relationship."
)

doc.add_heading("4.10.1 ERP Amplitudes by CDI", level=3)
doc.add_paragraph(
    "Mean FCz and Cz amplitudes by CDI condition:"
)

cdi_erp_rows = []
for cdi in ["rew_stay", "rew_shift", "norew_stay", "norew_shift"]:
    cdi_data = trial_df[trial_df["cond_dir"] == cdi]
    fcz_m = cdi_data["FCz"].mean()
    cz_m = cdi_data["Cz"].mean()
    n = len(cdi_data)
    cdi_erp_rows.append([cdi, f"{fcz_m:.4f}", f"{cz_m:.4f}", str(n)])
add_table(doc,
    ["CDI Condition", "FCz Mean (uV)", "Cz Mean (uV)", "N"],
    cdi_erp_rows
)

doc.add_paragraph("")

doc.add_heading("4.10.2 RPE-ERP Coupling by CDI", level=3)
doc.add_paragraph(
    "LME coefficients for RPE -> FCz within each CDI condition:"
)

cdi_rpe_rows = []
for rpe_type in ["rpe_binary", "rpe_mag"]:
    for cdi in ["rew_stay", "rew_shift", "norew_stay", "norew_shift"]:
        row = lme[(lme["predictor"] == rpe_type) & (lme["condition"] == cdi) &
                  (lme["channel"] == "FCz")]
        if len(row) > 0:
            r = row.iloc[0]
            sig = "*" if r["p"] < 0.05 else ""
            cdi_rpe_rows.append([
                rpe_type, cdi,
                f"{r['estimate']:.4f}", f"{r['se']:.4f}",
                f"{r['z']:.3f}", f"{r['p']:.4f}", sig
            ])
add_table(doc,
    ["Predictor", "CDI", "b", "SE", "z", "p", "Sig"],
    cdi_rpe_rows
)

doc.add_paragraph("")
doc.add_paragraph(
    "The CDI-specific analysis reveals condition-dependent patterns in "
    "RPE-ERP coupling. This is particularly informative because the CDI "
    "conditions decompose the overall effect into situations where the "
    "participant was rewarded vs. not, and stayed vs. shifted location."
)

doc.add_page_break()

# =====================================================================
# 5. FIGURE
# =====================================================================
doc.add_heading("5. Figure", level=1)

if os.path.exists(FIGURE_PATH):
    doc.add_picture(FIGURE_PATH, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run(
    "Figure 1. LME neural-behavioral fusion results for the T-maze AR task. "
    "(A) Best-fitting computational model per subject. "
    "(B) Binary RPE -> FCz LME coefficients by condition. "
    "(C) Magnitude RPE -> FCz LME coefficients by condition. "
    "(D) All predictor LME coefficients for FCz (all trials). "
    "(E) Binned binary RPE vs. FCz with LME statistics. "
    "(F) Binned magnitude RPE vs. FCz with LME statistics. "
    "(G) FCz amplitude by outcome x stay/shift. "
    "(H) Heatmap of binary RPE LME coefficients across channels and conditions. "
    "(I) RPE distribution by reward level. "
    "(J) Heatmap of magnitude RPE LME coefficients across channels and conditions. "
    "(K) Walking RT by reward level. "
    "(L) Learning rate distribution."
)
run.font.size = Pt(9)
run.italic = True

doc.add_paragraph("")

# Figure 2
doc.add_heading("5.2 CDI and RT_DIFF Analysis (Figure 2)", level=2)

if os.path.exists(FIGURE2_PATH):
    doc.add_picture(FIGURE2_PATH, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run(
    "Figure 2. CDI and RT_DIFF analyses. "
    "(A) FCz amplitude by CDI condition. "
    "(B) Cz amplitude by CDI condition. "
    "(C) RT_DIFF by CDI condition. "
    "(D) All predictors -> RT_DIFF LME coefficients. "
    "(E) Binary RPE -> FCz LME by CDI. "
    "(F) Magnitude RPE -> FCz LME by CDI. "
    "(G) Binary RPE -> RT_DIFF LME by CDI. "
    "(H) Magnitude RPE -> RT_DIFF LME by CDI. "
    "(I) RT_DIFF LME heatmap (predictors x conditions). "
    "(J) Binary RPE -> ERP heatmap (CDI conditions). "
    "(K) Magnitude RPE -> ERP heatmap (CDI conditions). "
    "(L) RT_DIFF distribution by CDI."
)
run.font.size = Pt(9)
run.italic = True

doc.add_page_break()

# =====================================================================
# 6. DISCUSSION
# =====================================================================
doc.add_heading("6. Discussion", level=1)

doc.add_heading("6.1 Magnitude RPE Drives ERP Effects", level=2)
doc.add_paragraph(
    "The most striking finding is the clear dissociation between binary and "
    "magnitude RPE. While binary RPE (derived from a simple reward/no-reward "
    "signal) showed almost no relationship with ERP amplitudes, magnitude RPE "
    "(derived from graded reward amounts: $0.00, $0.05, $0.25) was a robust "
    "predictor across multiple channels and conditions. This suggests that the "
    "brain's feedback processing system, as indexed by these ERP components, "
    "encodes the scalar magnitude of the prediction error rather than merely "
    "whether an outcome was rewarding or not."
)

doc.add_heading("6.2 Broad Scalp Distribution", level=2)
doc.add_paragraph(
    "Significant effects were not limited to the traditional frontocentral sites "
    "(FCz, Cz). Channels such as C4, P3, and P8 showed comparable numbers of "
    "significant effects. This broader topography may reflect the spatial "
    "navigation component of the AR T-maze task, which engages parietal and "
    "lateral neural systems in addition to the medial frontal generators of "
    "the classic RewP/FRN."
)

doc.add_heading("6.3 Behavioral Relevance of RPEs", level=2)
doc.add_paragraph(
    "The significant relationship between trial t-1 RPE and trial t walking "
    "speed change demonstrates that model-derived prediction errors have "
    "behavioral consequences. Participants walked faster after rewarding "
    "outcomes (behavioral invigoration) and slower after low-reward outcomes, "
    "consistent with the role of dopaminergic RPE signals in motivating action."
)

doc.add_heading("6.4 LME vs. Correlation Approach", level=2)
doc.add_paragraph(
    "The LME approach revealed significant effects that were not detected "
    "by the previous within-subject correlation approach. In particular, the "
    "magnitude RPE effect at FCz (z = -6.11) was highly significant in the LME "
    "framework. The increased power comes from: (1) using all single trials "
    "rather than summarizing per subject, (2) properly modeling the hierarchical "
    "data structure, and (3) accounting for trial-position effects through the "
    "variance component."
)

doc.add_heading("6.5 CDI-Specific Patterns", level=2)
doc.add_paragraph(
    "The Condition_Direction_index analysis reveals that the RPE-ERP "
    "relationship is not uniform across experimental conditions. By crossing "
    "reward outcome (reward vs. no-reward) with behavioral strategy (stay vs. "
    "shift), the CDI conditions isolate situations where prediction errors have "
    "different behavioral meanings. For instance, RPEs during rew_stay trials "
    "occur when a participant's repeated strategy is confirmed, while RPEs "
    "during norew_shift trials occur when exploration fails. These contextual "
    "differences may modulate the neural response to prediction errors."
)

doc.add_heading("6.6 RT_DIFF as a Behavioral Signature of Learning", level=2)
doc.add_paragraph(
    "The RT_DIFF analysis provides converging evidence that model-derived RPEs "
    "have real behavioral consequences. Magnitude RPE was the strongest "
    "behavioral predictor of RT_DIFF (after controlling for Start_RT regression "
    "to the mean), with negative coefficients indicating that larger positive "
    "RPEs drive speeding up on the next trial. This is consistent with "
    "dopaminergic RPE signals facilitating vigor and approach behavior. "
    "The fact that 64 out of 101 RT_DIFF models survived FDR correction "
    "suggests robust behavioral modulation by reward-learning variables."
)

doc.add_heading("6.7 Limitations", level=2)
doc.add_paragraph(
    "Several limitations should be noted. First, some models with the "
    "(1 | Trial) variance component showed convergence warnings (non-positive "
    "definite Hessian), indicating potential overparameterization. Second, "
    "the ERP amplitudes in this dataset represent a single time-window summary "
    "rather than a full time-resolved analysis. Third, FDR correction was "
    "applied across all 793 tests simultaneously; alternative correction "
    "strategies (e.g., separate correction within each predictor or channel "
    "family) could yield different significance patterns."
)

doc.add_page_break()

# =====================================================================
# 7. METHODS SUMMARY
# =====================================================================
doc.add_heading("7. Methods Summary (for Paper)", level=1)

doc.add_paragraph(
    "Single-trial ERP amplitudes from 13 scalp channels were predicted by "
    "trial-level behavioral variables using linear mixed-effects models. "
    "For each combination of behavioral predictor (11 predictors), EEG channel "
    "(13 channels), and condition subset (11 conditions including 4 CDI "
    "conditions crossing outcome x stay/shift), we fit the model: "
    "ERP ~ Predictor + (1 | Subject) + (1 | Trial), where (1 | Subject) captures "
    "between-subject baseline differences and (1 | Trial) captures shared "
    "trial-position effects via a variance component. Additionally, RT_DIFF "
    "(previous - current trial walking RT) was modeled as an outcome: "
    "RT_DIFF ~ Predictor + (1 | Subject) + (1 | Trial). "
    "Models were estimated with REML using the Powell optimizer in Python "
    "statsmodels 0.14.6. Multiple comparison correction used the "
    "Benjamini-Hochberg FDR procedure at q = .05, applied separately for "
    "ERP and RT_DIFF analyses. Interaction models tested whether RPE-ERP "
    "and RPE-RT_DIFF relationships were moderated by reward level, "
    "stay/shift behavior, velocity condition, and CDI."
)

# =====================================================================
# 8. OUTPUT FILES
# =====================================================================
doc.add_heading("8. Output Files", level=1)

add_table(doc,
    ["File", "Description", "Rows"],
    [
        ["lme_all_results.csv", f"All {len(lme)} ERP LME model results", f"{len(lme)}"],
        ["rtdiff_lme_results.csv", f"All {len(rtdiff)} RT_DIFF LME model results", f"{len(rtdiff)}"],
        ["trial_data_with_rpes.csv", "Trial-level data with predictors, ERPs, RT_DIFF", "7,963"],
        ["subject_info.csv", "Per-subject model fitting summary", "43"],
        ["interaction_models.csv", "Interaction model coefficients (incl. CDI)", f"{len(inter)}"],
        ["12_lme_tmaze_pipeline.png", "12-panel ERP summary figure", "-"],
        ["12_lme_cdi_rtdiff.png", "12-panel CDI & RT_DIFF figure", "-"],
    ]
)

# =====================================================================
# SAVE
# =====================================================================
output_path = os.path.join(RESULTS_DIR, "LME_TMaze_Analysis_Report.docx")
doc.save(output_path)
print(f"Report saved: {output_path}")
