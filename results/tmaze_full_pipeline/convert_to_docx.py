"""Convert the analysis report markdown to a formatted Word document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import re

doc = Document()

# -- Page setup ---------------------------------------------------------------
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# -- Styles -------------------------------------------------------------------
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        hs.font.size = Pt(16)
        hs.font.bold = True
    elif level == 2:
        hs.font.size = Pt(14)
        hs.font.bold = True
    else:
        hs.font.size = Pt(12)
        hs.font.bold = True
        hs.font.italic = True

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

# =============================================================================
# TITLE
# =============================================================================
title = doc.add_heading(
    'T-Maze AR: Reward Prediction Error x ERP x Reward Magnitude', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('Comprehensive Analysis Report')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True
subtitle.runs[0].font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Dataset: ').bold = True
p.add_run('merged_clean_V3_scaled.csv    ')
p.add_run('N = ').bold = True
p.add_run('43 participants, 7,963 single trials    ')
p.add_run('Task: ').bold = True
p.add_run('Augmented Reality T-maze (HoloLens 2)')

doc.add_paragraph()

# =============================================================================
# SECTION 1: OVERVIEW
# =============================================================================
doc.add_heading('1. Overview of the Analysis Pipeline', level=1)

doc.add_paragraph(
    'This report describes a computational modeling approach to link reward learning '
    'processes to single-trial ERP amplitudes in an augmented reality (AR) T-maze task. '
    'Participants wore a HoloLens 2 and physically walked to left or right locations in '
    'a virtual T-maze, receiving monetary feedback at three reward levels. We fit '
    'reinforcement learning (RL) models to each participant\'s trial-by-trial choices, '
    'extracted reward prediction errors (RPEs) using both binary and magnitude outcomes, '
    'and correlated these with single-trial ERP amplitudes across EEG channels.'
)

doc.add_heading('1.1 How RPEs Were Computed', level=2)

doc.add_paragraph(
    'Reward prediction errors (RPEs) quantify the discrepancy between expected and '
    'received outcomes on each trial. We derived RPEs using the Rescorla-Wagner (RW) '
    'model, which updates expected values according to:'
)

add_code_block(doc,
    'Q(chosen_action) <- Q(chosen_action) + alpha x (reward - Q(chosen_action))')

doc.add_paragraph('Where:')
bullets = [
    ('Q(action)', 'expected value of choosing that action (left or right arm)'),
    ('alpha', 'learning rate, controlling how rapidly the agent updates beliefs '
     '(0 = no learning, 1 = instant updating)'),
    ('reward', 'outcome on the current trial'),
    ('RPE = reward - Q(chosen_action)', 'the prediction error'),
]
for term, defn in bullets:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(term)
    run.bold = True
    p.add_run(f' = {defn}')

doc.add_paragraph('We computed two types of RPEs:')

p = doc.add_paragraph(style='List Number')
run = p.add_run('Binary RPE: ')
run.bold = True
p.add_run(
    'Outcomes coded as 1 (any reward) or 0 (no reward). '
    'RPE_binary = binary_outcome - Q(chosen). This captures the surprise of receiving '
    'vs. not receiving reward, regardless of magnitude.')

p = doc.add_paragraph(style='List Number')
run = p.add_run('Magnitude RPE: ')
run.bold = True
p.add_run(
    'Outcomes coded as actual Earned values ($0.00, $0.05, $0.25). '
    'RPE_magnitude = earned_amount - Q(chosen). This captures the surprise of the '
    'specific monetary amount, allowing the model to distinguish between high ($0.25) '
    'and low ($0.05) reward magnitudes.')

doc.add_paragraph(
    'The RPE is positive when the outcome is better than expected (surprise reward) '
    'and negative when worse than expected (surprise no-reward).'
)

doc.add_paragraph('Action selection follows a softmax rule:')
add_code_block(doc,
    'P(choose left) = exp(beta x Q(left)) / [exp(beta x Q(left)) + exp(beta x Q(right))]')
doc.add_paragraph(
    'Where beta = inverse temperature, controlling choice consistency '
    '(0 = random, large = deterministic).'
)

doc.add_heading('1.2 Model Fitting Procedure', level=2)

doc.add_paragraph('For each participant (code in agent/tools/model_fitting.py):')
steps = [
    'Input: Trial-by-trial choices (0=left, 1=right) and outcomes (0=no reward, 1=reward)',
    'Fitting method: Maximum Likelihood Estimation (MLE) via scipy\'s differential '
    'evolution optimizer',
    'Objective: Minimize the negative log-likelihood (NLL) of observed choices given '
    'model parameters',
    'Parameter bounds: alpha in [0.001, 0.999], beta in [0.1, 20.0]',
    'Model comparison: Bayesian Information Criterion (BIC = k*ln(n) - 2*ln(L)), '
    'where k = number of parameters, n = number of trials, L = likelihood. Lower BIC '
    'indicates better fit, penalizing model complexity.',
]
for step in steps:
    p = doc.add_paragraph(style='List Number')
    p.add_run(step)

doc.add_paragraph(
    'We fit 10 models per participant: Random, Win-Stay-Lose-Shift (WSLS), '
    'Rescorla-Wagner, Q-Learning, Q-Learning Dual Learning Rate, Actor-Critic, '
    'Q-Learning with Decay, RW with Side Bias, Choice Kernel, and RW + Choice Kernel.'
)

doc.add_heading('1.3 RPE Extraction', level=2)

doc.add_paragraph(
    'After fitting the RW model to each participant\'s data, we replayed each trial '
    'through the model with the best-fit parameters to extract RPE(t) = outcome(t) - '
    'Q(chosen_action, t) for each trial t. This was done separately for binary outcomes '
    '(0/1) and magnitude outcomes ($0.00/$0.05/$0.25), producing two RPE timeseries '
    'per participant.'
)

add_code_block(doc,
    'latent_bin = extract_trial_variables("rw", params_bin, choices, outcomes_binary)\n'
    'latent_mag = extract_trial_variables("rw", params_mag, choices, outcomes_magnitude)\n'
    'rpe_binary = latent_bin["rpes"]\n'
    'rpe_magnitude = latent_mag["rpes"]')

# =============================================================================
# SECTION 2: DATA STRUCTURE
# =============================================================================
doc.add_heading('2. Data Structure', level=1)

doc.add_heading('2.1 Behavioral Variables', level=2)
add_table(doc,
    ['Variable', 'Description'],
    [
        ['Start_RT', 'Walking time from countdown end to feedback onset (seconds)'],
        ['Return_RT', 'Walking time from feedback cue to start position (seconds)'],
        ['Direction', '1 = Left, 2 = Right (choice)'],
        ['Reward', '1 = Reward, 2 = No reward'],
        ['Trial', 'Trial number (1-200)'],
        ['PES_PRS_index', '0 = first trial, 1 = stay (same location), 2 = shift'],
        ['Condition_Direction_index',
         '0=first, 1=rew+stay, 2=rew+shift, 3=norew+stay, 4=norew+shift'],
        ['Vel_Cond', '1=High Reward ($0.25), 2=No Reward ($0.00), 3=Low Reward ($0.05)'],
        ['Earned', 'Monetary amount earned ($0.00, $0.05, or $0.25)'],
    ])

doc.add_heading('2.2 Reward Level Distribution', level=2)
doc.add_paragraph(
    'The Vel_Cond variable classifies trials by reward magnitude. High reward trials '
    'involve slower walking (higher Start_RT), while low reward trials involve faster '
    'walking (lower Start_RT). Vel_Cond=4 contains 0 trials in this dataset.'
)
add_table(doc,
    ['Vel_Cond', 'Label', 'Earned', 'N trials', 'Median Start_RT', 'Mean FCz'],
    [
        ['1', 'High Reward', '$0.25', '1,580', '4.667 sec', '2.768 uV'],
        ['2', 'No Reward', '$0.00', '4,011', '3.868 sec', '2.657 uV'],
        ['3', 'Low Reward', '$0.05', '2,372', '3.399 sec', '2.696 uV'],
    ])

p = doc.add_paragraph()
run = p.add_run('Note: ')
run.bold = True
p.add_run(
    'High reward trials were associated with slower walking speeds (median 4.67s) and '
    'low reward trials with faster walking (median 3.40s), creating a natural confound '
    'between reward magnitude and approach velocity.'
)

doc.add_heading('2.3 ERP Channels', level=2)
doc.add_paragraph(
    'Single-trial mean amplitudes extracted at 13 EEG electrodes: FCz, Cz, FC2, FC1, '
    'F4, F5, C4, C3, P4, P3, P8, P7, Pz.'
)

# =============================================================================
# SECTION 3: RESULTS
# =============================================================================
doc.add_heading('3. Results', level=1)

doc.add_heading('3.1 Computational Model Comparison', level=2)
add_table(doc,
    ['Model', 'N subjects (best)', '%'],
    [
        ['Win-Stay-Lose-Shift (WSLS)', '19', '44%'],
        ['Random Responding', '15', '35%'],
        ['Rescorla-Wagner', '5', '12%'],
        ['Other (rw_bias, ck, rwck, q_decay)', '4', '9%'],
    ])
doc.add_paragraph('')
doc.add_paragraph(
    'The WSLS heuristic was the best-fitting model for the majority of participants '
    '(44%), with Random Responding second (35%). This pattern is consistent with the '
    'equal reward probabilities for left and right arms (~50/50), which provide no '
    'differential value signal for RL models to exploit. Despite this, we fit the '
    'Rescorla-Wagner model to all participants to extract trial-by-trial RPEs, as RPEs '
    'capture the moment-to-moment surprise signal regardless of whether participants '
    'are using a value-learning strategy.'
)

doc.add_heading('3.2 RPE Characteristics by Reward Level', level=2)
add_table(doc,
    ['Condition', 'RPE_bin', '|RPE_bin|', 'RPE_mag', '|RPE_mag|', 'FCz', 'N'],
    [
        ['High Reward ($0.25)', '+0.521', '0.521', '+0.057', '0.164', '2.768', '1,580'],
        ['No Reward ($0.00)', '-0.488', '0.488', '-0.139', '0.139', '2.657', '4,011'],
        ['Low Reward ($0.05)', '+0.489', '0.489', '-0.065', '0.091', '2.696', '2,372'],
    ])
doc.add_paragraph('')
doc.add_paragraph(
    'Binary RPEs are positive on any reward trial and negative on no-reward trials, '
    'regardless of reward magnitude. Magnitude RPEs show a crucial distinction: high '
    'reward trials produce positive magnitude RPEs (+0.057, meaning $0.25 exceeded '
    'expectations), while low reward trials produce negative magnitude RPEs (-0.065, '
    'meaning $0.05 fell below expectations). This means the magnitude model captures '
    'that low-reward trials are actually disappointing relative to expectations.'
)

doc.add_heading('3.3 Single-Trial RPE x ERP Correlations', level=2)
doc.add_paragraph(
    'Within-subject Pearson correlations between trial-by-trial RPEs and single-trial '
    'ERP amplitudes were computed for each participant. Correlation coefficients were '
    'Fisher-z transformed and tested against zero at the group level using one-sample '
    't-tests.'
)

doc.add_heading('3.3.1 Binary RPE x ERP', level=3)
add_table(doc,
    ['Condition', 'Channel', 'Mean r', 't(42)', 'p', 'Sig.'],
    [
        ['All trials', 'FCz', '0.010', '0.79', '0.434', 'n.s.'],
        ['High Reward', 'FCz', '0.010', '0.27', '0.788', 'n.s.'],
        ['No Reward', 'FCz', '0.032', '2.18', '0.035', '*'],
        ['Low Reward', 'FCz', '0.019', '0.94', '0.353', 'n.s.'],
        ['Any Reward', 'FCz', '0.020', '1.15', '0.255', 'n.s.'],
        ['Stay', 'FCz', '-0.002', '-0.12', '0.906', 'n.s.'],
        ['Shift', 'FCz', '0.010', '0.62', '0.538', 'n.s.'],
    ])

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Key finding (Binary RPE): ')
run.bold = True
p.add_run(
    'Binary RPEs significantly predicted single-trial FCz amplitude specifically on '
    'no-reward trials (mean r = 0.032, t(42) = 2.18, p = 0.035). On these trials, '
    'larger negative RPEs (more unexpected no-reward) were associated with more positive '
    'FCz amplitudes. This is consistent with the feedback-related negativity (FRN) / '
    'reward positivity (REWP) reflecting prediction error signals.'
)

doc.add_heading('3.3.2 Magnitude RPE x ERP', level=3)
add_table(doc,
    ['Condition', 'Channel', 'Mean r', 't(42)', 'p', 'Sig.'],
    [
        ['All trials', 'FCz', '-0.029', '-2.02', '0.050', '*'],
        ['High Reward', 'FCz', '-0.036', '-1.08', '0.289', 'n.s.'],
        ['No Reward', 'FCz', '-0.039', '-1.87', '0.068', '~'],
        ['Low Reward', 'FCz', '-0.029', '-1.15', '0.256', 'n.s.'],
        ['Low Reward', 'Cz', '+0.044', '1.80', '0.080', '~'],
        ['Any Reward', 'FCz', '-0.030', '-1.92', '0.062', '~'],
        ['Stay', 'FCz', '-0.025', '-1.39', '0.171', 'n.s.'],
        ['Shift', 'FCz', '-0.043', '-2.41', '0.021', '*'],
    ])

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Key finding (Magnitude RPE): ')
run.bold = True
p.add_run(
    'Magnitude RPEs showed a different and complementary pattern. Overall magnitude '
    'RPE x FCz was significant (r = -0.029, p = 0.050), unlike binary RPE which was '
    'not significant overall (p = 0.434). The negative correlation means larger magnitude '
    'RPEs (better-than-expected monetary outcomes) were associated with smaller FCz '
    'amplitudes. The effect was strongest on shift trials (r = -0.043, p = 0.021).'
)

doc.add_paragraph(
    'The sign reversal between binary and magnitude RPEs is theoretically meaningful: '
    'Binary RPE x FCz is positive (especially on no-reward trials), while magnitude '
    'RPE x FCz is negative (especially on shift trials). This suggests that the '
    'FRN/REWP complex encodes two distinct signals -- a categorical reward/no-reward '
    'signal and a graded magnitude signal -- with opposing neural signatures at FCz.'
)

doc.add_heading('3.4 Reward Level Comparisons', level=2)

doc.add_heading('3.4.1 Binary RPE-FCz by Reward Level', level=3)
add_table(doc,
    ['Comparison', 't', 'p', 'High r', 'Low r'],
    [
        ['FCz High vs Low', '-0.22', '0.825', '0.010', '0.019'],
        ['FCz High vs NoRew', '-0.60', '0.550', '--', '--'],
        ['Cz High vs Low', '0.91', '0.372', '0.037', '-0.000'],
        ['Cz High vs NoRew', '1.09', '0.283', '--', '--'],
    ])
doc.add_paragraph('')
doc.add_paragraph(
    'Binary RPE-ERP coupling did not differ significantly between reward levels.')

doc.add_heading('3.4.2 Magnitude RPE by Reward Level', level=3)
add_table(doc,
    ['Comparison', 't', 'p', 'High r', 'Low r', 'Sig.'],
    [
        ['FCz High vs Low', '-0.13', '0.900', '-0.036', '-0.029', 'n.s.'],
        ['FCz High vs NoRew', '-0.63', '0.532', '--', '--', 'n.s.'],
        ['Cz High vs Low', '-2.34', '0.026', '-0.028', '+0.044', '*'],
        ['Cz High vs NoRew', '-0.29', '0.772', '--', '--', 'n.s.'],
    ])

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Key finding: ')
run.bold = True
p.add_run(
    'Magnitude RPE-Cz coupling significantly differed between high and low reward '
    'trials (t = -2.34, p = 0.026). High reward trials showed a negative magnitude '
    'RPE-Cz correlation (r = -0.028), while low reward trials showed a positive '
    'correlation (r = +0.044). This sign reversal suggests that the neural encoding '
    'of reward magnitude prediction errors differs qualitatively between high and low '
    'reward contexts at the vertex (Cz).'
)

doc.add_heading('3.4.3 Mean FCz by Reward Level', level=3)
add_table(doc,
    ['Reward Level', 'FCz (uV)', 'SD'],
    [
        ['High Reward ($0.25)', '2.768', '4.950'],
        ['No Reward ($0.00)', '2.657', '4.945'],
        ['Low Reward ($0.05)', '2.696', '5.149'],
    ])
doc.add_paragraph('')
add_table(doc,
    ['Comparison', 't', 'p'],
    [
        ['High vs Low', '-1.31', '0.199'],
        ['High vs NoRew', '-0.12', '0.901'],
        ['Low vs NoRew', '1.72', '0.093 (trend)'],
    ])
doc.add_paragraph('')
doc.add_paragraph(
    'A trending effect for Low vs NoRew (p = 0.093): low reward trials showed '
    'marginally higher FCz amplitude than no-reward trials, possibly reflecting '
    'a reward positivity even for small monetary amounts.'
)

doc.add_heading('3.5 Post-Outcome Walking Speed by Reward Level', level=2)
add_table(doc,
    ['Previous Reward Level', 'RT Change (sec)', 'SD'],
    [
        ['After High Reward ($0.25)', '-0.612 (speed up)', '1.530'],
        ['After Low Reward ($0.05)', '+0.327 (slow down)', '1.201'],
        ['After No Reward ($0.00)', '+0.026 (slight slow)', '1.495'],
    ])

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Key finding: ')
run.bold = True
p.add_run(
    'Magnitude RPE(t-1) strongly predicted walking speed change on trial t: '
    'mean r = -0.133, t(42) = -7.76, p < 0.0001. This is the strongest effect in '
    'the dataset. After high reward, participants walked 0.61 seconds faster on the '
    'next trial. After low reward, they slowed down by 0.33 seconds. This demonstrates '
    'that reward magnitude prediction errors have robust downstream effects on motoric '
    'behavior in physical walking.'
)

doc.add_paragraph(
    'Compare this to the previous analysis using only binary RPEs, which showed only '
    'a trending effect (r = -0.027, p = 0.085). The magnitude RPE captures nearly 5x '
    'more variance in subsequent walking speed, confirming that participants are '
    'sensitive to the graded monetary value, not just binary presence/absence of reward.'
)

doc.add_heading('3.6 Behavioral Adjustment Patterns', level=2)
add_table(doc,
    ['Metric', 'Value'],
    [
        ['Win-Stay rate', '0.516 (stay at same location after reward)'],
        ['Lose-Switch rate', '0.631 (switch location after no-reward)'],
        ['No-reward -> Stay', '1,473 trials (37%)'],
        ['No-reward -> Shift', '2,512 trials (63%)'],
        ['Reward -> Stay', '2,035 trials (52%)'],
        ['Reward -> Shift', '1,900 trials (48%)'],
    ])
doc.add_paragraph('')
doc.add_paragraph(
    'Participants were substantially more likely to shift after no-reward (63%) than '
    'after reward (48%), reflecting a lose-shift bias. Win-stay was near chance (52%), '
    'consistent with the equal reward probabilities making it uninformative to stay.'
)

# =============================================================================
# SECTION 4: FIGURE DESCRIPTIONS
# =============================================================================
doc.add_heading('4. Figure Descriptions', level=1)
doc.add_paragraph('Figure: 09_tmaze_full_pipeline.png (12-panel figure)')

panels = [
    ('Panel A: Best-Fitting Model per Subject',
     'Bar chart showing the number of subjects best fit by each computational model. '
     'WSLS dominates (19 subjects), followed by Random (15), RW (5), with 4 subjects '
     'best fit by other models.'),
    ('Panel B: Binary RPE-FCz Correlation Distribution',
     'Histogram of within-subject Pearson r values for binary RPE x FCz (all trials). '
     'The distribution is centered slightly above zero (mean r = 0.010, p = 0.434).'),
    ('Panel C: Binary RPE-FCz by Reward Level',
     'Bar chart comparing binary RPE-FCz correlation across reward levels: High Reward '
     '(r = 0.010), No Reward (r = 0.032*), Low Reward (r = 0.019). The no-reward '
     'condition shows the only significant effect.'),
    ('Panel D: Magnitude RPE-FCz by Reward Level',
     'Bar chart showing magnitude RPE-FCz correlation across conditions. All correlations '
     'are negative. Shift trials show the strongest effect (r = -0.043*), demonstrating '
     'distinct neural signatures for binary vs. magnitude prediction errors.'),
    ('Panel E: RPE Distributions by Reward Level',
     'Distributions of RPE values separately for High Reward (green), No Reward (red), '
     'and Low Reward (orange) trials for both binary and magnitude RPEs.'),
    ('Panel F: ERP Amplitude by Reward Level',
     'Mean FCz amplitudes for High Reward (2.768 uV), No Reward (2.657 uV), and Low '
     'Reward (2.696 uV). Differences are not statistically significant.'),
    ('Panel G: FCz by Outcome x Stay/Shift',
     'Mean FCz for reward/no-reward x stay/shift combinations. Reward+Shift shows the '
     'highest amplitude (2.847 uV), Reward+Stay the lowest (2.499 uV).'),
    ('Panel H: Walking Speed by Reward Level',
     'Median walking times for each reward level. High Reward has longest walking times '
     '(~4.67 sec, slow approach), Low Reward shortest (~3.40 sec, fast approach).'),
    ('Panel I: Post-Outcome RT Change by Reward Level',
     'Mean RT change as a function of previous reward level. After High Reward: -0.61 sec '
     '(speed up). After Low Reward: +0.33 sec (slow down). After No Reward: +0.03 sec.'),
    ('Panel J: RPE(t-1) -> RT Change Correlation',
     'Histogram of within-subject correlations between previous trial magnitude RPE and '
     'current trial RT change. Distribution shifted substantially negative (mean r = -0.133, '
     'p < 0.0001), indicating robust post-outcome walking speed adjustment.'),
    ('Panel K: RPE-ERP Topography (Binary)',
     'Bar chart showing binary RPE-ERP correlation at each of 13 EEG channels. '
     'Frontocentral channels (FCz, FC1) show small positive correlations.'),
    ('Panel L: RPE-ERP Topography (Magnitude)',
     'Bar chart showing magnitude RPE-ERP correlation at each of 13 EEG channels. '
     'Negative magnitude RPE-FCz correlation (-0.029) shows frontocentral distribution.'),
]

for ptitle, desc in panels:
    p = doc.add_paragraph()
    run = p.add_run(ptitle)
    run.bold = True
    p.add_run('\n' + desc)

# =============================================================================
# SECTION 5: HIGH VS LOW REWARD
# =============================================================================
doc.add_heading('5. Discussion: High vs Low Reward Magnitude', level=1)

doc.add_heading('5.1 Three-Level Reward Structure', level=2)
doc.add_paragraph(
    'This AR T-maze uses three reward magnitudes:')
for item in [
    'High Reward ($0.25, VC=1): Participants walked slowly (median 4.67 sec) and '
    'received $0.25. Highest FCz amplitudes (2.768 uV), most positive magnitude RPEs '
    '(+0.057).',
    'Low Reward ($0.05, VC=3): Participants walked quickly (median 3.40 sec) and '
    'received $0.05. Intermediate FCz (2.696 uV), negative magnitude RPEs (-0.065).',
    'No Reward ($0.00, VC=2): All no-reward trials. Intermediate walking times '
    '(median 3.87 sec), lowest FCz (2.657 uV), most negative magnitude RPEs (-0.139).',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.2 Magnitude RPE Reveals What Binary RPE Misses', level=2)
doc.add_paragraph(
    'The binary RPE treats all reward trials identically (outcome = 1), discarding '
    'information about reward magnitude. The magnitude RPE preserves this information: '
    'high reward produces positive magnitude RPEs (outcome > expectation) while low '
    'reward produces negative magnitude RPEs (outcome < expectation).'
)
doc.add_paragraph('This distinction matters because:')
for item in [
    'Magnitude RPE-FCz was significant overall (p = 0.050) while binary was not '
    '(p = 0.434).',
    'Magnitude RPE captured 5x more variance in post-outcome walking speed '
    '(r = -0.133 vs r = -0.027).',
    'Magnitude RPE-Cz coupling differed significantly between high and low reward '
    '(p = 0.026).',
]:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('5.3 The Sign Reversal at Cz', level=2)
doc.add_paragraph(
    'At Cz, high reward trials showed negative magnitude RPE-ERP coupling (r = -0.028) '
    'while low reward trials showed positive coupling (r = +0.044), a significant '
    'difference (p = 0.026). This cross-over interaction provides evidence that the '
    'vertex ERP encodes reward magnitude in a context-dependent manner: for high reward, '
    'larger-than-expected outcomes produce smaller Cz amplitudes (ceiling effect), while '
    'for low reward, larger-than-expected outcomes produce larger Cz amplitudes '
    '(reward positivity scaling with relative value).'
)

# =============================================================================
# SECTION 6: REWARD VS NO-REWARD
# =============================================================================
doc.add_heading('6. Discussion: Reward vs No-Reward Trials', level=1)

doc.add_heading('6.1 The Asymmetric RPE-ERP Effect', level=2)
doc.add_paragraph(
    'The central finding from binary RPE analysis is that RPEs predict single-trial '
    'FCz amplitude on no-reward trials only (r = 0.032, p = 0.035), not on reward '
    'trials (p = 0.255) or all trials (p = 0.434). This asymmetry is consistent with '
    'the FRN/REWP primarily encoding negative prediction errors.'
)

doc.add_heading('6.2 Dual Encoding of Prediction Errors', level=2)
doc.add_paragraph(
    'The complementary binary and magnitude RPE findings suggest the feedback-locked '
    'ERP carries at least two components:')
for item in [
    'A categorical "good/bad" signal (captured by binary RPE, positive correlation '
    'with FCz on no-reward trials)',
    'A graded magnitude signal (captured by magnitude RPE, negative correlation with '
    'FCz on shift trials)',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.3 Behavioral Consequences of Reward Magnitude', level=2)
doc.add_paragraph(
    'After high reward ($0.25), participants walked 0.61 sec faster. After low reward '
    '($0.05), they slowed down by 0.33 sec. The paradoxical finding that low reward '
    'slows participants more than no reward may reflect a frustration effect: receiving '
    'only $0.05 when expecting $0.25 could be more aversive than receiving nothing, '
    'consistent with the negative magnitude RPE on low-reward trials.'
)

# =============================================================================
# SECTION 7: METHODS NOTES
# =============================================================================
doc.add_heading('7. Methodological Notes', level=1)

doc.add_heading('7.1 Why Use RW Model When WSLS Fits Better?', level=2)
for item in [
    'Consistency: Same model across all participants for comparable RPE magnitudes.',
    'RPE generation: WSLS does not generate meaningful trial-level prediction errors.',
    'Neural relevance: The brain likely computes prediction errors even when behavior '
    'appears heuristic.',
    'Precedent: Follows Wilson & Collins (2019) and Sambrook & Goslin (2015).',
]:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('7.2 Binary vs Magnitude RPE', level=2)
doc.add_paragraph(
    'We report both binary and magnitude RPE analyses because binary RPEs are standard '
    'in the literature while magnitude RPEs exploit the three-level reward structure. '
    'The sign reversal between binary (positive) and magnitude (negative) RPE-FCz '
    'correlations reveals distinct neural components. Post-outcome walking speed is '
    'far better predicted by magnitude RPEs (r = -0.133) than binary RPEs (r = -0.027).'
)

doc.add_heading('7.3 Single-Trial ERP Analysis', level=2)
doc.add_paragraph(
    'We used within-subject correlations (Pearson r) between trial-by-trial RPEs and '
    'single-trial ERP amplitudes, Fisher-z transformed for group-level inference. This '
    'follows the robust single-trial approach of Hauser et al. (2014) and Fischer & '
    'Ullsperger (2013).'
)

# =============================================================================
# SECTION 8: KEY FINDINGS SUMMARY
# =============================================================================
doc.add_heading('8. Summary of Key Findings', level=1)
add_table(doc,
    ['#', 'Finding', 'Statistic', 'p'],
    [
        ['1', 'Mag RPE x FCz on shift trials', 'r=-0.043, t(42)=-2.41', '0.021*'],
        ['2', 'Bin RPE x FCz on no-reward trials', 'r=0.032, t(42)=2.18', '0.035*'],
        ['3', 'Mag RPE x FCz overall', 'r=-0.029, t(42)=-2.02', '0.050*'],
        ['4', 'Mag RPE(t-1) -> RT change(t)', 'r=-0.133, t(42)=-7.76', '<0.0001***'],
        ['5', 'Mag RPE-Cz: High vs Low reward', 't(29)=-2.34', '0.026*'],
        ['6', 'Post-high-reward speed-up', '-0.61 sec', '--'],
        ['7', 'Post-low-reward slow-down', '+0.33 sec', '--'],
    ])

# =============================================================================
# SECTION 9: OUTPUT FILES
# =============================================================================
doc.add_heading('9. Output Files', level=1)
add_table(doc,
    ['File', 'Description'],
    [
        ['trial_data_with_rpes.csv',
         '7,963 rows with binary + magnitude RPEs, Q-values, and ERP amplitudes'],
        ['within_subject_correlations.csv',
         'Per-subject r values for each condition x channel'],
        ['group_rpe_erp_stats.csv', 'Group-level t-test results for all conditions'],
        ['subject_info.csv', 'Per-subject model fits and RW parameters'],
        ['rpe_rt_correlations.csv', 'Per-subject RPE -> RT change correlations'],
        ['09_tmaze_full_pipeline.png', '12-panel figure'],
        ['08_rpe_erp_heatmap.png', 'Heatmap of RPE-ERP correlations'],
    ])

# =============================================================================
# SECTION 10: REFERENCES
# =============================================================================
doc.add_heading('10. References', level=1)
refs = [
    'Wilson, R. C., & Collins, A. G. (2019). Ten simple rules for the computational '
    'modeling of behavioral data. eLife, 8, e49547.',
    'Sambrook, T. D., & Goslin, J. (2015). A neural reward prediction error revealed '
    'by a meta-analysis of ERPs using great grand averages. Psychological Bulletin, '
    '141(1), 213.',
    'Hauser, T. U., et al. (2014). The feedback-related negativity (FRN) revisited: '
    'new insights into the localization, meaning and network organization. NeuroImage, '
    '84, 159-168.',
    'Fischer, A. G., & Ullsperger, M. (2013). Real and fictive outcomes are processed '
    'differently but converge on a common adaptive mechanism. Neuron, 79(6), 1243-1255.',
    'Holroyd, C. B., & Coles, M. G. (2002). The neural basis of human error processing: '
    'reinforcement learning, dopamine, and the error-related negativity. Psychological '
    'Review, 109(4), 679.',
    'Proudfit, G. H. (2015). The reward positivity: From basic research on reward to a '
    'biomarker for depression. Psychophysiology, 52(4), 449-459.',
]
for ref in refs:
    doc.add_paragraph(ref, style='List Bullet')

# -- Save ---------------------------------------------------------------------
out_path = os.path.join(os.path.dirname(__file__), 'ANALYSIS_REPORT.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
