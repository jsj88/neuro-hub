from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading("Reward Learning Models — Master Catalog", 0)
doc.add_paragraph("Compiled: February 2026 | Sources: Wilson & Collins (2019) eLife; Ahn et al. (2014); Haines et al. (2018)")

# ── SECTION 1: Core Bandit Models (Wilson & Collins 2019) ──────────────────
doc.add_heading("1. Core Bandit Models (Wilson & Collins, 2019)", level=1)
doc.add_paragraph("Task: 2-armed bandit. All models use softmax: p(k) = exp(V_k) / Σ exp(V_i)")

models = [
    ("M1 — Random Responding", "b (bias)", "Null/Baseline",
     "p(k) = b  |  p(other) = 1 - b",
     "Participants who do not engage with reward structure. Captures fixed response bias."),
    ("M2 — Noisy Win-Stay-Lose-Shift (nWSLS)", "ε (noise)", "Heuristic RL",
     "p(k) = 1 - ε/2 if win-stay or lose-shift condition, else ε/2",
     "Repeat rewarded actions, switch from unrewarded. Stochastic version of a simple heuristic."),
    ("M3 — Rescorla-Wagner (RW)", "α (learning rate), β (inv. temperature)", "Model-free RL",
     "Q(k) ← Q(k) + α(r - Q(k))  |  p(k) = exp(βQ(k)) / Σ exp(βQ(i))",
     "Workhorse delta-rule RL. Tracks expected reward value via prediction errors."),
    ("M4 — Choice Kernel (CK)", "αc, βc", "Perseveration",
     "CK(k) ← CK(k) + αc(a(k) - CK(k))  |  p(k) = exp(βc·CK(k)) / Σ exp(βc·CK(i))",
     "Repeats recently chosen actions independent of reward. Captures perseveration/side biases."),
    ("M5 — RW + Choice Kernel (RW+CK)", "α, β, αc, βc", "RL + Perseveration",
     "V(k) = β·Q(k) + βc·CK(k)  |  Q and CK updated as M3 and M4",
     "Combines reward learning and action repetition tendency. Most complex bandit model."),
    ("M6 — RW + Side Bias", "α, β, B", "RL + Nuisance",
     "p(left) = 1 / (1 + exp(β(Q(right) - Q(left) - B)))",
     "Adds spatial/motor bias to RW. Including nuisance params improves α, β recovery."),
]

for name, params, category, equations, description in models:
    doc.add_heading(name, level=2)
    t = doc.add_table(rows=4, cols=2)
    t.style = "Table Grid"
    rows = [("Category", category), ("Parameters", params), ("Equations", equations), ("What it models", description)]
    for i, (label, val) in enumerate(rows):
        t.rows[i].cells[0].text = label
        t.rows[i].cells[1].text = val
    doc.add_paragraph("")

# ── SECTION 2: Extended Models ─────────────────────────────────────────────
doc.add_heading("2. Extended / State-Based Models", level=1)

extended = [
    ("Mixed RL + Working Memory (RL+WM)", "α, β, ρ (WM weight), K (WM capacity)",
     "RL + Working Memory", "Collins & Frank (2012), Eur J Neurosci",
     "w_WM = min(1, K/n_s) · ρ  |  Final policy = w_WM · WM + (1-w_WM) · RL",
     "Dissociates WM-driven one-shot learning from incremental RL. Stimulus-action learning tasks."),
    ("State-Based RL (fullRL)", "αP, αN, β",
     "Model-free RL with state", "—",
     "Q(a,s) ← Q(a,s) + αP/N · (r - Q(a,s))  |  separate rates for + and - PEs",
     "Learns separate action values per stimulus. Benchmark 'correct' model for S-A tasks."),
    ("Stimulus-Blind RL (blindRL)", "αP, αN, β",
     "Model-free RL (no state)", "—",
     "Q(a) ← Q(a) + αP/N · (r - Q(a))  (ignores stimulus identity)",
     "Wrong model for S-A tasks. Used to illustrate model validation failures."),
]

for name, params, category, citation, equations, description in extended:
    doc.add_heading(name, level=2)
    t = doc.add_table(rows=5, cols=2)
    t.style = "Table Grid"
    rows = [("Category", category), ("Citation", citation), ("Parameters", params), ("Equations", equations), ("What it models", description)]
    for i, (label, val) in enumerate(rows):
        t.rows[i].cells[0].text = label
        t.rows[i].cells[1].text = val
    doc.add_paragraph("")

# ── SECTION 3: Iowa Gambling Task Models ──────────────────────────────────
doc.add_heading("3. Iowa Gambling Task (IGT) Models", level=1)

doc.add_heading("Value-Plus-Perseverance (VPP) — Ahn et al. (2014)", level=2)
doc.add_paragraph("Citation: Ahn et al. (2014), Frontiers in Psychology, 5:849")
doc.add_paragraph("Category: Hybrid Model-free RL + Outcome-Modulated Perseverance")
doc.add_paragraph("Task: Iowa Gambling Task (4 decks, 100 trials)")
doc.add_paragraph("Parameters (8 total):")
vpp_params = [
    ("A", "(0,1)", "Learning rate — controls how strongly recent outcomes update deck expectancy"),
    ("α", "(0,2)", "Outcome sensitivity — curvature of prospect utility function"),
    ("λ", "(0,10)", "Loss aversion — scales sensitivity to losses vs. gains"),
    ("c", "(0,5)", "Response consistency → θ = 3^c - 1 (softmax temperature)"),
    ("εp", "(-∞,∞)", "Gain impact on perseverance (positive = stay after win)"),
    ("εn", "(-∞,∞)", "Loss impact on perseverance (negative = switch after loss)"),
    ("K", "(0,1)", "Perseverance decay across all decks"),
    ("ω", "(0,1)", "Weight on RL (E) vs. perseverance (P) — ω=1: pure RL; ω=0: pure heuristic"),
]
t = doc.add_table(rows=1 + len(vpp_params), cols=3)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Symbol"; t.rows[0].cells[1].text = "Bounds"; t.rows[0].cells[2].text = "Role"
for i, (sym, bounds, role) in enumerate(vpp_params):
    t.rows[i+1].cells[0].text = sym; t.rows[i+1].cells[1].text = bounds; t.rows[i+1].cells[2].text = role
doc.add_paragraph("\nEquations:")
for eq in [
    "Utility:  u(t) = x^α  if x≥0  ;  u(t) = -λ|x|^α  if x<0",
    "Value update (chosen):  E_j(t+1) = E_j(t) + A[u(t) - E_j(t)]",
    "Perseverance (chosen, gain):  P_j(t+1) = K·P_j(t) + εp",
    "Perseverance (chosen, loss):  P_j(t+1) = K·P_j(t) + εn",
    "Perseverance (unchosen):  P_k(t+1) = K·P_k(t)",
    "Combined value:  V_j = ω·E_j + (1-ω)·P_j",
    "Choice:  p(j) = exp(θ·V_j) / Σ exp(θ·V_k)  where θ = 3^c - 1",
]:
    doc.add_paragraph(eq, style="List Bullet")
doc.add_paragraph("")

doc.add_heading("Outcome Representation Learning (ORL) — Haines et al. (2018)", level=2)
doc.add_paragraph("Citation: Haines, Vassileva & Ahn (2018), Cognitive Science, 42(8):2534–2561")
doc.add_paragraph("Category: Model-free RL with separate value + frequency representations")
doc.add_paragraph("Task: Iowa Gambling Task (4 decks)")
doc.add_paragraph("Parameters (5 total):")
orl_params = [
    ("Arew", "(0,1)", "Reward learning rate — updates after gains"),
    ("Apun", "(0,1)", "Punishment learning rate — updates after losses"),
    ("K", "(0,5)", "Perseverance decay (K_tr = 3^K - 1)"),
    ("βF", "(-∞,∞)", "Outcome frequency weight — scales ef contribution to utility"),
    ("βP", "(-∞,∞)", "Perseverance weight — positive=stay, negative=switch"),
]
t = doc.add_table(rows=1 + len(orl_params), cols=3)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Symbol"; t.rows[0].cells[1].text = "Bounds"; t.rows[0].cells[2].text = "Role"
for i, (sym, bounds, role) in enumerate(orl_params):
    t.rows[i+1].cells[0].text = sym; t.rows[i+1].cells[1].text = bounds; t.rows[i+1].cells[2].text = role
doc.add_paragraph("\nEquations:")
for eq in [
    "Value PE:         δ_val(t) = o(t) - ev_j(t)",
    "Frequency PE:     δ_freq(t) = sign(o(t)) - ef_j(t)",
    "Fictive freq PE:  δ_fic_k(t) = -sign(o(t))/3 - ef_k(t)  [unchosen decks]",
    "Value update (gain):  ev_j(t+1) = ev_j(t) + Arew·δ_val",
    "Value update (loss):  ev_j(t+1) = ev_j(t) + Apun·δ_val",
    "Freq update: chosen with Arew/Apun; unchosen via fictive PE with swapped LRs",
    "Perseverance:  pers_j(t+1) = (pers_j(t) + 1[j=chosen]) / (1 + K_tr)",
    "Utility:  util_j = ev_j + βF·ef_j + βP·pers_j",
    "Choice:  p(j) = exp(util_j) / Σ exp(util_k)",
]:
    doc.add_paragraph(eq, style="List Bullet")
doc.add_paragraph("")

# ── SECTION 4: VPP vs ORL Comparison ──────────────────────────────────────
doc.add_heading("4. VPP vs. ORL — Direct Comparison", level=1)
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
headers = ["Feature", "VPP (Ahn 2014)", "ORL (Haines 2018)"]
for i, h in enumerate(headers):
    t.rows[0].cells[i].text = h
comparisons = [
    ("Parameters", "8", "5"),
    ("Learning rates", "Single (A)", "Asymmetric (Arew, Apun)"),
    ("Loss aversion", "Explicit λ", "Implicit via LR asymmetry"),
    ("Outcome representation", "Single utility (prospect fn)", "Separate value (ev) + frequency (ef)"),
    ("Fictive updating", "No", "Yes — unchosen decks updated"),
    ("Perseverance", "Outcome-modulated (εp, εn)", "Binary reset, exponential decay"),
    ("Utility combination", "Weighted avg: ω·E + (1-ω)·P", "Additive: ev + βF·ef + βP·pers"),
    ("Choice temp parameter", "Explicit c → θ", "Absorbed into β weights"),
    ("hBayesDM function", "igt_vpp()", "igt_orl()"),
]
for row in comparisons:
    cells = t.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# ── SECTION 5: Referenced Models ──────────────────────────────────────────
doc.add_heading("5. Referenced / Cited Models (not fully implemented in TenSimpleRules repo)", level=1)
ref_models = [
    ("Q-Learning / TD Learning", "Model-free RL", "Watkins & Dayan (1992)"),
    ("Temporal Difference (TD) Learning", "Model-free RL", "Montague, Dayan & Sejnowski (1996)"),
    ("Model-Based RL (two-step task)", "Model-based RL", "Daw et al. (2011), Neuron"),
    ("Prospect Theory", "Utility-based", "Kahneman & Tversky"),
    ("Cumulative Prospect Theory (CPT)", "Utility-based", "Nilsson et al. (2011)"),
    ("Ideal Observer Model", "Bayesian/normative", "Geisler (2011)"),
    ("OpAL (Opponent Actor Learning)", "Actor-critic RL", "Collins & Frank (2014)"),
    ("Approx. Bayesian Delta Rule", "Bayesian RL", "Nassar et al. (2010)"),
    ("Mixture of Delta-Rules", "Bayesian RL approx.", "Wilson, Nassar & Gold (2013)"),
    ("Bayesian Bandit Analysis", "Full Bayesian RL", "Steyvers, Lee & Wagenmakers (2009)"),
    ("Particle Filter", "Bayesian sequential", "Courville & Daw (2008)"),
    ("Drift-Diffusion Model (DDM)", "Sequential sampling/RT", "Ratcliff (1978)"),
    ("Feature-Based Learning", "Attentional RL", "Farashahi et al. (2017)"),
    ("Attention + RL (RLWM)", "Attentional RL", "Leong et al. (2017), Neuron"),
    ("Resource-Rational Models", "Bounded rationality", "Lieder et al. (2018)"),
    ("Pavlovian-Instrumental Interaction", "Pavlovian + instrumental RL", "Huys et al. (2011)"),
    ("Dopaminergic RL in Parkinsonism", "Asymmetric RL", "Frank et al. (2004), Science"),
    ("Bayesian Adaptive Delta Rule", "Bayesian RL", "Nassar et al. (2012), Nat Neurosci"),
]
t = doc.add_table(rows=1 + len(ref_models), cols=3)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Model"; t.rows[0].cells[1].text = "Category"; t.rows[0].cells[2].text = "Citation"
for i, row in enumerate(ref_models):
    for j, val in enumerate(row):
        t.rows[i+1].cells[j].text = val

# ── Save ───────────────────────────────────────────────────────────────────
output_path = "/Users/jaleesastringfellow/neuro-hub/reward_learning_models.docx"
doc.save(output_path)
print(f"Saved: {output_path}")